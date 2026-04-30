#!/usr/bin/env python3
"""Build Word deliverable from static template + results paths (SP-GQE)."""

from __future__ import annotations

from pathlib import Path

REPO = Path(__file__).resolve().parents[1]
OUT = REPO / "deliverables" / "SP_GQE_Empirical_Report.docx"
RES = REPO / "results"


def main() -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    d = Document()
    s = d.sections[0]
    s.left_margin = Inches(1)
    s.right_margin = Inches(1)

    t = d.add_paragraph("SP-GQE: Empirical evaluation on HotpotQA (distractor)")
    t.runs[0].font.size = Pt(16)
    t.runs[0].bold = True
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph("Robert — dissertation deliverable (April 2026)").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    d.add_paragraph()

    p = d.add_paragraph(
        "Abstract. We report a multi-seed aggregate of SP-GQE and baselines (V-RAG, "
        "GQE-RAG, GR-RAG, GF-RAG) on 562 HotpotQA distractor dev instances, 23 random seeds. "
        "Paired 95% bootstrap CIs for (SP-GQE−V-RAG) token F1 include zero on bridge and "
        "comparison subsets. Graph-query validity ablations show complementary two-branch "
        "pools and a precision–recall trade-off under cosine pruning. Full detail: see "
        "companion Markdown."
    )
    p.runs[0].italic = True

    d.add_heading("1. Introduction and motivation", level=1)
    d.add_paragraph(
        "Open-domain multi-hop question answering (Yang et al., 2018) and retrieval-augmented "
        "generation (Lewis et al., 2020) are standard frameworks. This work evaluates a "
        "query-time, training-free design (SPARQL on a per-question co-occurrence graph + "
        "semantic pruning + dense retrieval + Groq reader) with pipeline controls held fixed "
        "as far as the protocol allows (config/EXPERIMENT_PROTOCOL.md). Motivation: test whether "
        "graph-conditioned queries improve token F1 vs classical dense RAG, especially on "
        "bridge questions (pre-registered H1)."
    )

    d.add_heading("2. Methods (summary)", level=1)
    d.add_paragraph(
        "Pipelines: V-RAG; GQE-RAG(n=2); SP-GQE(n=2,τ=0.5); SP-GQE-i; GR-RAG; GF-RAG. "
        "Reader: llama-3.1-8b-instant (T=0) under --stack plan. "
        "Aggregation: mean ± 95% CI on seed means; paired ΔF1 with bootstrap 95% CI. "
        "Re-run: python scripts/aggregate_daily_runs.py"
    )

    d.add_heading("3. Results — tables", level=1)
    d.add_paragraph("Table 1. Per-pipeline (mean of seed means, 23 seeds; 95% CI on F1):")
    tbl1 = d.add_table(rows=7, cols=6)
    hdr = [
        "Pipeline",
        "Mean F1",
        "95% CI F1",
        "Mean EM",
        "Sup. recall@k",
        "P@k",
    ]
    for c, t in enumerate(hdr):
        tbl1.rows[0].cells[c].text = t
    data = [
        (
            "GQE-RAG(n=2)",
            "0.5760",
            "[0.5381, 0.6139]",
            "0.4636",
            "0.8021",
            "0.6654",
        ),
        (
            "GR-RAG",
            "0.5727",
            "[0.5361, 0.6092]",
            "0.4607",
            "0.7908",
            "0.6738",
        ),
        (
            "V-RAG",
            "0.5633",
            "[0.5276, 0.5991]",
            "0.4520",
            "0.7908",
            "0.6738",
        ),
        (
            "SP-GQE-i",
            "0.5520",
            "[0.5212, 0.5828]",
            "0.4560",
            "0.7781",
            "0.6610",
        ),
        (
            "GF-RAG",
            "0.5512",
            "[0.5213, 0.5811]",
            "0.4403",
            "0.7445",
            "0.6382",
        ),
        (
            "SP-GQE(n=2,τ=0.5)",
            "0.5489",
            "[0.5149, 0.5828]",
            "0.4463",
            "0.7704",
            "0.6360",
        ),
    ]
    for r, row in enumerate(data, start=1):
        for c, val in enumerate(row):
            tbl1.rows[r].cells[c].text = val

    d.add_paragraph("Table 2. Paired ΔF1: SP-GQE(n=2,τ=0.5) − V-RAG (pooled):")
    tbl2 = d.add_table(rows=3, cols=4)
    for c, t in enumerate(["Subset", "Mean Δ", "Bootstrap 95% CI", "n pairs"]):
        tbl2.rows[0].cells[c].text = t
    tbl2.rows[1].cells[0].text = "bridge"
    tbl2.rows[1].cells[1].text = "-0.0185"
    tbl2.rows[1].cells[2].text = "[-0.0473, 0.0093]"
    tbl2.rows[1].cells[3].text = "270"
    tbl2.rows[2].cells[0].text = "comparison"
    tbl2.rows[2].cells[1].text = "-0.0096"
    tbl2.rows[2].cells[2].text = "[-0.0448, 0.0253]"
    tbl2.rows[2].cells[3].text = "292"

    d.add_paragraph("Table 3. Graph-query validity (pooled, n=562):")
    tbl3 = d.add_table(rows=5, cols=4)
    h3 = ["Stage", "Mean P", "Mean R", "n Q"]
    for c, t in enumerate(h3):
        tbl3.rows[0].cells[c].text = t
    gv = [
        ("Branch 1 n-hop", "0.3460", "0.5959", "562"),
        ("Branch 2 keyword", "0.3793", "0.1698", "562"),
        ("Union", "0.3155", "0.6343", "562"),
        ("Kept @τ=0.5", "0.4545", "0.2192", "562"),
    ]
    for r, row in enumerate(gv, start=1):
        for c, val in enumerate(row):
            tbl3.rows[r].cells[c].text = val

    d.add_heading("4. Figures", level=1)
    for name, cap in [
        (
            "pipelines_bar_f1.png",
            "Bar chart: mean F1 by pipeline (last local run; tables above are full aggregate).",
        ),
        ("heatmap_fungi_n_tau.png", "n×τ heatmap (mean F1) — single-seed sensitivity run; illustrative."),
        (
            "heatmap_fungi_n_tau_retrieval_p_at_k.png",
            "n×τ heatmap (mean P@k) — same.",
        ),
    ]:
        path = RES / name
        d.add_paragraph(f"Figure: {name}")
        if path.is_file():
            d.add_picture(str(path), width=Inches(5.8))
        else:
            d.add_paragraph(f"[Image not found: {path}]")
        d.add_paragraph(cap)
        d.add_paragraph()

    d.add_heading("5. References (selection)", level=1)
    refs = [
        "Lewis, P., et al. (2020). Retrieval-augmented generation for knowledge-intensive NLP. NeurIPS.",
        "Yang, Z., et al. (2018). HotpotQA: diverse, explainable multi-hop QA. EMNLP.",
        "Edge, D., et al. (2024). From local to global: graph RAG for query-focused summarization. (Microsoft / GraphRAG, technical).",
    ]
    for r in refs:
        d.add_paragraph(r, style="List Number")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
