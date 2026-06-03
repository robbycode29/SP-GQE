#!/usr/bin/env python3
"""Build Raport Cercetare 3 (semester preliminary) and 4 (full) as English DOCX."""

from __future__ import annotations

import json
import re
from pathlib import Path

REPO = Path(__file__).resolve().parents[1]
RES = REPO / "results"
OUT_DIR = REPO / "deliverables" / "rapoarte-cercetare"
SUMMARY_PATH = RES / "AGGREGATED_SUMMARY.json"
QUERIES_SAMPLE = RES / "daily_runs" / "2026-05-07__seed42__n25__queries.md"

# Semester snapshot (23 seeds, 562 instances) — before heatmap merge / validity reporting
SEMESTER = {
    "n_seeds": 23,
    "n_instances": 562,
    "pipelines": [
        ("GQE-RAG (n=2)", "0.5760", "[0.5381, 0.6139]", "0.4636", "0.8021", "0.6654"),
        ("GR-RAG", "0.5727", "[0.5361, 0.6092]", "0.4607", "0.7908", "0.6738"),
        ("V-RAG", "0.5633", "[0.5276, 0.5991]", "0.4520", "0.7908", "0.6738"),
        ("SP-GQE-i (n=3, τ=0.5)", "0.5520", "[0.5212, 0.5828]", "0.4560", "0.7781", "0.6610"),
        ("GF-RAG", "0.5512", "[0.5213, 0.5811]", "0.4403", "0.7445", "0.6382"),
        ("SP-GQE (n=2, τ=0.5)", "0.5489", "[0.5149, 0.5828]", "0.4463", "0.7704", "0.6360"),
    ],
    "paired_delta": [
        ("bridge", "-0.0185", "[-0.0473, 0.0093]", "270"),
        ("comparison", "-0.0096", "[-0.0448, 0.0253]", "292"),
    ],
}

SEED42_HEATMAP_F1 = [
    [0.506, 0.517, 0.520, 0.517, 0.426],
    [0.549, 0.517, 0.517, 0.519, 0.426],
    [0.592, 0.517, 0.520, 0.519, 0.426],
]
TAUS = ["0.3", "0.4", "0.5", "0.6", "0.7"]
N_HOPS = ["1", "2", "3"]


def _fmt_ci(ci: list[float] | tuple[float, float]) -> str:
    return f"[{ci[0]:.4f}, {ci[1]:.4f}]"


def _load_summary() -> dict:
    with SUMMARY_PATH.open(encoding="utf-8") as f:
        return json.load(f)


def _setup_doc(title: str, subtitle: str):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    d = Document()
    s = d.sections[0]
    s.left_margin = Inches(1)
    s.right_margin = Inches(1)
    t = d.add_paragraph(title)
    t.runs[0].font.size = Pt(16)
    t.runs[0].bold = True
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = d.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph()
    return d


def _add_abstract(d, text: str) -> None:
    p = d.add_paragraph(text)
    if p.runs:
        p.runs[0].italic = True


def _add_table(d, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    tbl = d.add_table(rows=1 + len(rows), cols=len(headers))
    for c, h in enumerate(headers):
        tbl.rows[0].cells[c].text = h
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            tbl.rows[r].cells[c].text = str(val)
    d.add_paragraph()


def _add_figure(d, path: Path, caption: str, width_in: float = 5.8) -> None:
    from docx.shared import Inches

    d.add_paragraph(f"Figure: {path.name}")
    if path.is_file():
        d.add_picture(str(path), width=Inches(width_in))
    else:
        d.add_paragraph(f"[Image not found: {path}]")
    d.add_paragraph(caption)
    d.add_paragraph()


def _add_code_block(d, code: str) -> None:
    from docx.shared import Pt

    for line in code.strip().splitlines():
        p = d.add_paragraph(line)
        if p.runs:
            p.runs[0].font.name = "Consolas"
            p.runs[0].font.size = Pt(9)


def _format_entity_list(raw: str, *, max_show: int = 10) -> str:
    """Compact display for returned-entity lists (handles truncated markdown)."""
    items = re.findall(r"'((?:\\'|[^'])*)'", raw)
    if not items:
        return raw.strip()[:350] + ("…" if len(raw) > 350 else "")
    if len(items) <= max_show:
        return "[" + ", ".join(f"'{x}'" for x in items) + "]"
    head = ", ".join(f"'{x}'" for x in items[:max_show])
    return f"[{head}, … (+{len(items) - max_show} more)]"


def _parse_returned(part: str) -> tuple[int | None, str]:
    m = re.search(r"- \*\*returned \(n=(\d+)\):\*\* (.+)", part)
    if not m:
        return None, ""
    return int(m.group(1)), _format_entity_list(m.group(2).strip())


def _parse_question_block(block: str) -> dict:
    m_title = re.match(r"\[(\w+)\] (.+?)\n", block)
    qtype = m_title.group(1) if m_title else "?"
    question = m_title.group(2).strip() if m_title else block[:80]

    def _meta(pat: str) -> str | None:
        m = re.search(pat, block)
        return m.group(1).strip() if m else None

    branches: list[dict] = []
    parts = re.split(r"### Branch \d+ — ", block)
    for part in parts[1:]:
        label_line = part.split("\n", 1)[0].strip()
        sq = re.search(r"```sparql\n(.*?)```", part, re.DOTALL)
        n_ret, entities = _parse_returned(part)
        branches.append(
            {
                "label": label_line,
                "sparql": sq.group(1).strip() if sq else "",
                "returned_n": n_ret,
                "returned_entities": entities,
            }
        )

    union_m = re.search(r"\*\*union \|A ∪ B\| = (\d+)\*\*", block)
    kept_m = re.search(
        r"\*\*kept after τ = [\d.]+ \(n = (\d+)\):\*\* (.+)", block
    )
    sims = re.findall(r"- `([^`]+)` — ([\d.]+)", block)
    top_sims = [(e, s) for e, s in sims[:8]]

    validity_rows: list[tuple[str, str, str]] = []
    for row in re.finditer(
        r"\| (Branch 1 \(n-hop\)|Branch 2 \(keyword\)|Union|Kept after τ) "
        r"\| ([\d.]+) \| ([\d.]+) \|",
        block,
    ):
        validity_rows.append((row.group(1), row.group(2), row.group(3)))

    supporting = _meta(r"\*\*supporting entities.*?:\*\* (.+)")
    if supporting and len(supporting) > 280:
        supporting = _format_entity_list(supporting, max_show=8)

    return {
        "type": qtype,
        "question": question,
        "qid": _meta(r"\*\*qid:\*\* `([^`]+)`"),
        "gold_answer": _meta(r"\*\*gold answer:\*\* (.+)"),
        "f1_sp_gqe": _meta(r"\*\*SP-GQE F1 / V-RAG F1:\*\* ([\d.]+)"),
        "f1_v_rag": _meta(r"\*\*SP-GQE F1 / V-RAG F1:\*\* [\d.]+ / ([\d.]+)"),
        "seed_entities": _meta(r"\*\*seed entities.*?:\*\* (.+)"),
        "branches": branches,
        "union_size": union_m.group(1) if union_m else None,
        "kept_n": kept_m.group(1) if kept_m else None,
        "kept_entities": (
            _format_entity_list(kept_m.group(2).strip(), max_show=12)
            if kept_m
            else ""
        ),
        "top_similarities": top_sims,
        "supporting_entities": supporting,
        "validity_rows": validity_rows,
    }


def _extract_sparql_samples(md_path: Path, max_questions: int = 3) -> list[dict]:
    text = md_path.read_text(encoding="utf-8")
    blocks = re.split(r"\n## \d+\. ", text)
    return [_parse_question_block(b) for b in blocks[1 : 1 + max_questions]]


def _add_sparql_sample_to_doc(d, samp: dict) -> None:
    title = f"[{samp['type']}] {samp['question'][:120]}"
    d.add_heading(title, level=2)
    meta = []
    if samp.get("qid"):
        meta.append(f"qid: {samp['qid']}")
    if samp.get("gold_answer"):
        meta.append(f"gold answer: {samp['gold_answer']}")
    if samp.get("f1_sp_gqe") and samp.get("f1_v_rag"):
        meta.append(
            f"answer F1 — SP-GQE: {samp['f1_sp_gqe']}, V-RAG: {samp['f1_v_rag']}"
        )
    if samp.get("seed_entities"):
        meta.append(f"seed entities: {samp['seed_entities']}")
    if meta:
        d.add_paragraph(" | ".join(meta))

    for br in samp["branches"]:
        d.add_paragraph(f"Branch — {br['label']}")
        if br.get("sparql"):
            _add_code_block(d, br["sparql"])
        if br.get("returned_n") is not None:
            d.add_paragraph(
                f"Query results (n={br['returned_n']}): {br['returned_entities']}"
            )

    d.add_paragraph("Fusion and pruning")
    if samp.get("union_size"):
        d.add_paragraph(f"Union |A ∪ B| = {samp['union_size']} entities.")
    if samp.get("kept_n"):
        d.add_paragraph(
            f"Kept after τ = 0.5 (n = {samp['kept_n']}): {samp['kept_entities']}"
        )
    if samp.get("top_similarities"):
        d.add_paragraph("Top candidate cosine similarities (entity → score):")
        for ent, score in samp["top_similarities"]:
            d.add_paragraph(f"  • {ent} — {score}", style="List Bullet")

    if samp.get("validity_rows"):
        d.add_paragraph("Graph-query validity (vs gold supporting NER):")
        _add_table(
            d,
            ["Stage", "Precision", "Recall"],
            [(a, b, c) for a, b, c in samp["validity_rows"]],
        )
    if samp.get("supporting_entities"):
        d.add_paragraph(
            f"Gold supporting entities (reference): {samp['supporting_entities']}"
        )
    d.add_paragraph()


def build_report_3() -> Path:
    d = _setup_doc(
        "Research Report 3 — SP-GQE empirical evaluation (semester preliminary)",
        "Robert Oprescu — past semester snapshot (English)",
    )
    _add_abstract(
        d,
        "This report summarises the cumulative multi-seed evaluation of SP-GQE and in-repo "
        "baselines on a stratified HotpotQA distractor subset (562 question instances, "
        "23 random seeds). The primary configuration SP-GQE (n=2, τ=0.5) is compared to "
        "V-RAG under a fixed protocol (same corpus, embeddings, and Groq reader). "
        "Preliminary sensitivity over hop count n and cosine threshold τ is reported from "
        "per-seed grid runs (illustrative single-seed heatmaps); multi-seed heatmap "
        "aggregation and graph-query validity ablations are deferred to Research Report 4.",
    )

    d.add_heading("1. Introduction", level=1)
    d.add_paragraph(
        "Multi-hop open-domain QA (Yang et al., 2018) and retrieval-augmented generation "
        "(Lewis et al., 2020) motivate a training-free pipeline that queries a per-question "
        "co-occurrence RDF graph with SPARQL, fuses two branches, cosine-prunes entity "
        "candidates, and augments dense retrieval before a small LLM reader. This semester "
        "report documents pipeline-level means and paired (SP-GQE − V-RAG) token F1 on "
        "bridge and comparison subsets, plus exploratory (n, τ) grids without merged "
        "cross-seed selection or validity-stage ablations."
    )

    d.add_heading("2. Methods (summary)", level=1)
    d.add_paragraph(
        "Pipelines: V-RAG; GQE-RAG (n=2); SP-GQE (n=2, τ=0.5); SP-GQE-i (n=3, τ=0.5); "
        "GR-RAG; GF-RAG. Reader: Groq llama-3.1-8b-instant (T=0). Aggregation: mean ± 95% "
        "CI on seed-level means; paired ΔF1 with bootstrap 95% CI. Protocol: "
        "config/EXPERIMENT_PROTOCOL.md."
    )

    d.add_heading("3. Results — pipeline comparison", level=1)
    d.add_paragraph(
        f"Table 1. Per-pipeline means ({SEMESTER['n_seeds']} seeds, "
        f"{SEMESTER['n_instances']} instances; 95% CI on F1 across seed means)."
    )
    _add_table(
        d,
        ["Pipeline", "Mean F1", "95% CI F1", "Mean EM", "Sup. recall@k", "P@k"],
        SEMESTER["pipelines"],
    )

    d.add_paragraph(
        "Table 2. Paired ΔF1: SP-GQE (n=2, τ=0.5) − V-RAG (question-level, pooled)."
    )
    _add_table(
        d,
        ["Subset", "Mean Δ", "Bootstrap 95% CI", "n pairs"],
        SEMESTER["paired_delta"],
    )

    d.add_heading("4. Preliminary (n, τ) sensitivity", level=1)
    d.add_paragraph(
        "Each heatmap-enabled seed run sweeps n ∈ {1, 2, 3} and τ ∈ {0.3, …, 0.7} for "
        "SP-GQE on the same sampled questions. Below: mean token F1 grid for seed 42 "
        "(n=25, 2026-05-07). These grids are exploratory; no sample-size-weighted merge "
        "across seeds is applied in this report."
    )
    heat_rows = []
    for i, nh in enumerate(N_HOPS):
        row = [nh] + [f"{SEED42_HEATMAP_F1[i][j]:.3f}" for j in range(5)]
        heat_rows.append(tuple(row))
    _add_table(d, ["n hops"] + [f"τ={t}" for t in TAUS], heat_rows)

    for name, cap in [
        (
            "heatmap_fungi_n_tau.png",
            "Single-seed mean answer F1 over (n, τ) — illustrative sensitivity.",
        ),
        (
            "heatmap_fungi_n_tau_retrieval_p_at_k.png",
            "Single-seed mean retrieval P@k over (n, τ) — same run.",
        ),
    ]:
        _add_figure(d, RES / name, cap)

    d.add_heading("5. Discussion and next steps", level=1)
    d.add_paragraph(
        "At the protocol default (n=2, τ=0.5), mean SP-GQE F1 does not exceed V-RAG; "
        "paired bootstrap CIs for ΔF1 include zero on bridge and comparison subsets. "
        "Exploratory grids suggest τ and n affect both F1 and P@k; a follow-up campaign "
        "will aggregate heatmaps across seeds, select a global (n*, τ*), and report "
        "graph-query validity (Research Report 4)."
    )

    d.add_heading("References (selection)", level=1)
    for ref in [
        "Lewis, P., et al. (2020). Retrieval-augmented generation for knowledge-intensive NLP. NeurIPS.",
        "Yang, Z., et al. (2018). HotpotQA: diverse, explainable multi-hop QA. EMNLP.",
        "Edge, D., et al. (2024). From local to global: Graph RAG (technical report / blog).",
    ]:
        d.add_paragraph(ref, style="List Number")

    out = OUT_DIR / "Oprescu Robert - Raport Cercetare 3.docx"
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    try:
        d.save(str(out))
    except PermissionError:
        alt = OUT_DIR / "Oprescu Robert - Raport Cercetare 3 (regenerated).docx"
        d.save(str(alt))
        print(
            f"WARNING: could not overwrite {out} (file open?). Wrote {alt} instead."
        )
        return alt
    return out


def build_report_4(summary: dict) -> Path:
    agg = summary["aggregated_across_seeds"]
    n_files = summary["n_files"]
    n_inst = summary["total_question_instances"]
    proto = summary["protocol_default_n2_tau05"]
    merged = summary["merged_config_aggregates"]
    sel = merged["selection"]
    sp_sel = merged["sp_gqe_at_selected_config"]
    gv = proto["graph_query_validity_pooled"]

    d = _setup_doc(
        "Research Report 4 — SP-GQE empirical evaluation (complete)",
        "Robert Oprescu — cumulative results through May 2026 (English)",
    )
    _add_abstract(
        d,
        f"This report is the complete empirical update: {n_inst} stratified HotpotQA "
        f"distractor instances across {n_files} seeds. It reports per-pipeline aggregates, "
        f"merged-grid selection of (n*, τ*) from {sel['merged_n_source_files']} heatmap "
        f"runs (weight {sel['merged_total_weight']}), SP-GQE at the selected cell, paired "
        f"ΔF1 versus V-RAG, graph-query validity at the protocol default, and qualitative "
        f"SPARQL samples. Limitations where logging does not match the selected (n, τ) are "
        f"stated explicitly.",
    )

    d.add_heading("1. Introduction", level=1)
    d.add_paragraph(
        "Building on Research Report 3, this document adds multi-seed heatmap aggregation, "
        "global (n, τ) selection by argmax on pooled mean F1, retrieval heatmaps, "
        "graph-query validity (precision/recall of entity sets vs gold supporting NER), "
        "and annotated SPARQL query samples from daily run logs."
    )

    d.add_heading("2. Evaluation scale", level=1)
    seeds = summary.get("seeds_used", [])
    d.add_paragraph(
        f"Aggregated from {n_files} per-seed JSON files ({n_inst} unique question instances). "
        f"Heatmap merge: {sel['merged_n_source_files']} files, total weight "
        f"{sel['merged_total_weight']}."
    )
    if seeds:
        seed_str = ", ".join(
            f"{s['date_utc']}:seed{s['seed']}(n={s['sample_size']})" for s in seeds[:8]
        )
        if len(seeds) > 8:
            seed_str += f", … (+{len(seeds) - 8} more)"
        d.add_paragraph(f"Seeds (sample): {seed_str}")

    d.add_heading("3. Per-pipeline results (protocol arms)", level=1)
    pipe_order = [
        "V-RAG",
        "GQE-RAG(n=2)",
        "SP-GQE(n=2,τ=0.5)",
        "SP-GQE-i(n=3,τ=0.5)",
        "GR-RAG",
        "GF-RAG",
    ]
    rows = []
    for name in pipe_order:
        p = agg[name]
        rows.append(
            (
                name,
                f"{p['mean_f1']:.4f}",
                _fmt_ci(p["ci95_f1"]),
                f"{p['mean_em']:.4f}",
                f"{p['mean_supporting_title_recall_at_k']:.4f}",
                f"{p['mean_retrieval_p_at_k']:.4f}",
                str(p["n_seeds"]),
            )
        )
    n_star, tau_star = int(sel["n_hops"]), float(sel["tau"])
    posthoc_label = f"SP-GQE(n={n_star},τ={tau_star})"
    rows.append(
        (
            posthoc_label,
            f"{sp_sel['mean_f1']:.4f}",
            _fmt_ci(sp_sel["ci95_f1"]),
            "—",
            "—",
            f"{sp_sel['mean_retrieval_p_at_k']:.4f}",
            str(sp_sel["n_seeds"]),
        )
    )
    _add_table(
        d,
        [
            "Pipeline",
            "Mean F1",
            "95% CI F1",
            "Mean EM",
            "Sup. recall@k",
            "P@k",
            "n seeds",
        ],
        rows,
    )
    d.add_paragraph(
        f"{posthoc_label} (last row): post-hoc argmax on pooled heatmap mean F1 "
        f"(see Section 4); not a pre-registered protocol arm. EM and supporting-title "
        f"recall were not swept in the (n, τ) grid. n={sp_sel['n_seeds']} heatmap "
        f"seeds (protocol arms: n={agg['V-RAG']['n_seeds']} seeds)."
    )

    d.add_heading("4. Heatmap aggregation and selected (n, τ)", level=1)
    d.add_paragraph(
        f"Selection criterion: {sel['criterion']}. "
        f"Global best on merged grid: n={sel['n_hops']}, τ={sel['tau']} "
        f"(pooled mean F1 = {sel['pooled_mean_f1_on_merged_grid']:.4f}, "
        f"P@k = {sel['pooled_mean_p_at_k_on_merged_grid']:.4f}). "
        f"Per-seed metrics use this same cell (not per-seed argmax)."
    )
    d.add_paragraph(
        f"SP-GQE (n={sel['n_hops']}, τ={sel['tau']}) — mean F1 "
        f"{sp_sel['mean_f1']:.4f} {_fmt_ci(sp_sel['ci95_f1'])}; "
        f"mean P@k {sp_sel['mean_retrieval_p_at_k']:.4f} "
        f"{_fmt_ci(sp_sel['ci95_retrieval_p_at_k'])}; n={sp_sel['n_seeds']} heatmap seeds."
    )
    if merged.get("sp_gqe_i_at_selected_config") is None:
        d.add_paragraph(
            "SP-GQE-i at the merged cell: not yet logged in older daily JSON "
            "(heatmap_sp_gqe_i added in newer run_experiment.py)."
        )

    d.add_paragraph("Figures — merged heatmaps (sample-size-weighted across heatmap seeds).")
    for path, cap in [
        (
            RES / "merged_heatmap" / "heatmap_merged_mean_f1.png",
            "Merged mean answer F1 over (n, τ); star marks pooled argmax.",
        ),
        (
            RES / "merged_heatmap" / "heatmap_merged_mean_p_at_k.png",
            "Merged mean retrieval P@k over (n, τ).",
        ),
    ]:
        _add_figure(d, path, cap)

    d.add_heading("5. Paired ΔF1 (SP-GQE − V-RAG)", level=1)
    pd_seed = merged["paired_delta_f1_seed_level_SP_GQE_minus_V_RAG"]
    d.add_paragraph("Table — seed-level ΔF1 at merged-selected (n*, τ*):")
    _add_table(
        d,
        ["Scope", "Mean Δ", "95% CI", "n"],
        [
            (
                f"all seeds (SP-GQE n={sel['n_hops']}, τ={sel['tau']})",
                f"{pd_seed['mean_delta_f1']:.4f}",
                _fmt_ci(pd_seed["ci95_delta_f1"]),
                str(pd_seed["n_seeds"]),
            ),
        ],
    )
    pd_q = merged["paired_delta_f1_question_level"]
    d.add_paragraph(
        f"Question-level pairs (heatmap seeds only; SP-GQE arm in log: "
        f"{pd_q['sp_gqe_arm_in_paired_deltas']}). Bridge/comparison rows do not use "
        f"the heatmap cell F1 unless re-logged per question."
    )
    _add_table(
        d,
        ["Subset", "Mean Δ", "Bootstrap 95% CI", "n pairs"],
        [
            (
                "bridge",
                f"{pd_q['bridge']['mean_delta_f1']:.4f}",
                _fmt_ci(pd_q["bridge"]["bootstrap_ci95"]),
                str(pd_q["bridge"]["n_pairs"]),
            ),
            (
                "comparison",
                f"{pd_q['comparison']['mean_delta_f1']:.4f}",
                _fmt_ci(pd_q["comparison"]["bootstrap_ci95"]),
                str(pd_q["comparison"]["n_pairs"]),
            ),
        ],
    )
    d.add_paragraph("Appendix — protocol default SP-GQE (n=2, τ=0.5):")
    pd_proto = proto["paired_delta_f1_SP_GQE_minus_V_RAG"]
    _add_table(
        d,
        ["Subset", "Mean Δ", "Bootstrap 95% CI", "n pairs"],
        [
            (
                "bridge",
                f"{pd_proto['bridge']['mean']:.4f}",
                _fmt_ci(pd_proto["bridge"]["bootstrap_ci95"]),
                str(pd_proto["bridge"]["n_pairs"]),
            ),
            (
                "comparison",
                f"{pd_proto['comparison']['mean']:.4f}",
                _fmt_ci(pd_proto["comparison"]["bootstrap_ci95"]),
                str(pd_proto["comparison"]["n_pairs"]),
            ),
        ],
    )

    d.add_heading("6. Graph-query validity", level=1)
    d.add_paragraph(
        merged.get(
            "graph_query_validity_note",
            "Validity at merged-selected (n, τ) was not logged; table uses protocol default.",
        )
    )
    nq = gv["n_questions"]
    _add_table(
        d,
        ["Stage", "Mean precision", "Mean recall", "n questions"],
        [
            (
                "Branch 1 (n-hop)",
                f"{gv['branch1_nhop']['mean_precision']:.4f}",
                f"{gv['branch1_nhop']['mean_recall']:.4f}",
                str(nq),
            ),
            (
                "Branch 2 (keyword)",
                f"{gv['branch2_keyword']['mean_precision']:.4f}",
                f"{gv['branch2_keyword']['mean_recall']:.4f}",
                str(nq),
            ),
            (
                "Union",
                f"{gv['union']['mean_precision']:.4f}",
                f"{gv['union']['mean_recall']:.4f}",
                str(nq),
            ),
            (
                f"Kept after τ=0.5",
                f"{gv['kept_after_tau']['mean_precision']:.4f}",
                f"{gv['kept_after_tau']['mean_recall']:.4f}",
                str(nq),
            ),
        ],
    )

    d.add_heading("7. SPARQL query samples", level=1)
    d.add_paragraph(
        f"Excerpted from {QUERIES_SAMPLE.name} (seed 42, SP-GQE n=2, τ=0.5). "
        "For each question we show both branch queries, entity sets returned by each "
        "SPARQL execution, fusion/pruning outcomes, top cosine scores, and per-stage "
        "precision/recall against gold supporting entities."
    )
    if QUERIES_SAMPLE.is_file():
        for samp in _extract_sparql_samples(QUERIES_SAMPLE, max_questions=3):
            _add_sparql_sample_to_doc(d, samp)
    else:
        d.add_paragraph(f"[Queries file not found: {QUERIES_SAMPLE}]")

    d.add_heading("8. Conclusions", level=1)
    d.add_paragraph(
        f"Across {n_files} seeds, GQE-RAG and GR-RAG remain competitive on mean F1; "
        f"merged-grid SP-GQE at (n={sel['n_hops']}, τ={sel['tau']}) reaches "
        f"{sp_sel['mean_f1']:.4f} mean F1 on heatmap seeds but paired improvements "
        f"over V-RAG are not statistically decisive at α≈0.05 (CIs include zero). "
        f"Validity diagnostics show complementary branches (union recall > branch 2) "
        f"and higher precision after τ pruning at the cost of recall."
    )

    d.add_heading("References (selection)", level=1)
    for ref in [
        "Lewis, P., et al. (2020). Retrieval-augmented generation for knowledge-intensive NLP. NeurIPS.",
        "Yang, Z., et al. (2018). HotpotQA: diverse, explainable multi-hop QA. EMNLP.",
        "Edge, D., et al. (2024). From local to global: Graph RAG (technical report / blog).",
    ]:
        d.add_paragraph(ref, style="List Number")

    out = OUT_DIR / "Oprescu Robert - Raport Cercetare 4.docx"
    try:
        d.save(str(out))
    except PermissionError:
        alt = OUT_DIR / "Oprescu Robert - Raport Cercetare 4 (regenerated).docx"
        d.save(str(alt))
        print(
            f"WARNING: could not overwrite {out} (file open?). Wrote {alt} instead."
        )
        return alt
    return out


def _save_doc(d, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        d.save(str(path))
        return path
    except PermissionError:
        alt = path.with_stem(path.stem + " (regenerated)")
        d.save(str(alt))
        print(f"WARNING: could not overwrite {path} (file open?). Wrote {alt} instead.")
        return alt


def build_motivation_method_docx() -> Path:
    d = _setup_doc(
        "SP-GQE: Motivation and Method",
        "Robert Oprescu — dissertation design document (English)",
    )
    _add_abstract(
        d,
        "SP-GQE (Semantic Pruning on Graph-structured Query Expansion) is a training-free "
        "retrieval-augmented QA pipeline for HotpotQA distractor multi-hop questions. A "
        "per-question RDF co-occurrence graph is queried with two complementary SPARQL 1.1 "
        "branches; entity candidates are fused and cosine-pruned before augmenting a shared "
        "dense retriever and reader. This document states the research motivation, "
        "pre-registered hypothesis, related-work positioning, and the implemented method "
        "(graph construction, SPARQL branches, pruning, baselines, metrics, and statistical "
        "design). Empirical outcomes are reported separately in Research Reports 3–4.",
    )

    d.add_heading("1. Research motivation", level=1)
    d.add_paragraph(
        "Multi-hop open-domain question answering requires locating and combining evidence "
        "across several passages (Yang et al., 2018). Retrieval-augmented generation (RAG) "
        "(Lewis et al., 2020) remains the standard recipe: embed the question, retrieve "
        "top-k paragraphs with a dense index, and answer with a language model reader. "
        "When a question hinges on linking entities or relations that are not all lexical "
        "neighbours of the surface question, dense retrieval alone may miss bridging facts "
        "unless the retrieval query encodes intermediate structure."
    )
    d.add_paragraph(
        "Graph-augmented systems attach textual or relational structure at index time, "
        "query time, or both—community summaries, entity graphs, subgraph retrieval, and "
        "hybrid dense–graph rankers (e.g. Edge et al., 2024; HybGRAG; SubgraphRAG). A "
        "design question for dissertation-scale work is whether lightweight, query-time "
        "expansion on a graph built only from the provided corpus can improve answer "
        "quality under strict comparability: same paragraphs, same embedding model, same "
        "FAISS index, and same reader—varying only how the retrieval query is conditioned."
    )
    d.add_paragraph(
        "SP-GQE targets that question with an explicit, inspectable mechanism: SPARQL over "
        "a per-question RDF graph, two branches that capture structural co-occurrence and "
        "label-level keyword overlap, and a semantic pruning step that filters noisy entity "
        "candidates before they augment dense retrieval. The approach is training-free, uses "
        "no external knowledge base, and is bounded to HotpotQA distractor dev paragraphs "
        "(ten passages per question)."
    )

    d.add_heading("2. Research questions and hypothesis", level=1)
    d.add_paragraph(
        "Primary research question: Does graph-structured query expansion with semantic "
        "pruning improve answer token F1 over classical dense RAG when all other stack "
        "components are held fixed?"
    )
    d.add_paragraph(
        "Pre-registered hypothesis (H1, config/EXPERIMENT_PROTOCOL.md): SP-GQE improves "
        "token F1 versus V-RAG on bridge-type HotpotQA items, where multi-hop linking of "
        "evidence is central. Null expectation: on comparison-type items (typically "
        "two-entity comparisons), dense retrieval often suffices and SP-GQE is not expected "
        "to outperform V-RAG reliably."
    )
    d.add_paragraph(
        "Secondary diagnostic questions: (i) Do the two SPARQL branches return complementary "
        "entity sets (union recall exceeding either branch alone)? (ii) Does cosine pruning "
        "at threshold τ trade precision for recall in a way that explains downstream retrieval "
        "and answer behaviour? These are evaluated via graph-query validity ablations."
    )

    d.add_heading("3. Related work (positioning)", level=1)
    _add_table(
        d,
        ["Theme", "Representative work", "Relation to SP-GQE"],
        [
            (
                "Multi-hop QA benchmark",
                "Yang et al. (2018) HotpotQA",
                "Distractor dev; stratified bridge / comparison sampling.",
            ),
            (
                "Dense RAG",
                "Lewis et al. (2020)",
                "V-RAG baseline: FAISS on question only, same reader.",
            ),
            (
                "Document / community graphs",
                "Edge et al. (2024) GraphRAG",
                "Different scope: in-corpus entity co-occurrence, not global summaries.",
            ),
            (
                "Dense + graph hybrids",
                "HybGRAG, SubgraphRAG (cited, not re-implemented)",
                "Motivates fair same-reader pipeline comparisons.",
            ),
        ],
    )

    d.add_heading("4. System overview", level=1)
    d.add_paragraph(
        "End-to-end flow for each question q:"
    )
    for step in [
        "Build RDF graph G(q) from the ten distractor paragraphs (entities + sentence co-occurrence).",
        "Extract seed entities (spaCy NER on q) and noun-chunk probes (≤12, length-filtered).",
        "Branch 1: SPARQL n-hop traversal on spg:coOccurs from seeds → set A.",
        "Branch 2: SPARQL keyword filter on rdfs:label from probes → set B.",
        "Fusion: candidates = (A ∪ B) \\ seeds; score each by max cosine to {q} ∪ probes; keep if ≥ τ; always retain seeds.",
        "Dense retrieval: FAISS top-k with augmented query string (graph context + probes).",
        "Reader: Groq llama-3.1-8b-instant, temperature 0, short answer.",
    ]:
        d.add_paragraph(step, style="List Number")

    d.add_heading("5. Graph construction and SPARQL branches", level=1)
    d.add_heading("5.1 Per-question RDF graph", level=2)
    d.add_paragraph(
        "For each question, an RDF graph is materialised in-memory (rdflib, RdfQuestionGraph) "
        "from the HotpotQA distractor paragraph set only:"
    )
    for item in [
        "Triple: ⟨e⟩ rdf:type spg:Entity for every spaCy-NER entity normalised to a canonical label.",
        "Triple: ⟨e⟩ rdfs:label \"normalised surface form\".",
        "Symmetric edge: ⟨a⟩ spg:coOccurs ⟨b⟩ when a and b co-occur in the same sentence (both directions stored).",
    ]:
        d.add_paragraph(item, style="List Bullet")

    d.add_heading("5.2 Branch 1 — structural n-hop (SPARQL 1.1)", level=2)
    d.add_paragraph(
        "Seed entities are spaCy-NER mentions in the question (normalised). A bounded property-path "
        "query unions paths of length 1…n over spg:coOccurs, e.g. for n=2:"
    )
    _add_code_block(
        d,
        """PREFIX spg: <http://spgqe.local/>
SELECT DISTINCT ?t WHERE {
  VALUES ?s { ... seed URIs ... }
  { ?s spg:coOccurs ?t } UNION
  { ?s spg:coOccurs/spg:coOccurs ?t }
}""",
    )
    d.add_paragraph(
        "The result set A is the set of entity labels reachable within n hops. This branch "
        "exploits local graph topology and is intended to surface bridge entities that co-occur "
        "with question mentions even when they are not direct string matches."
    )

    d.add_heading("5.3 Branch 2 — keyword semantic SPARQL", level=2)
    d.add_paragraph(
        "Keywords are derived from spaCy noun chunks on the question (lowercased, deduplicated, "
        "stop-word filtered, minimum length 3). A second query filters entities by substring "
        "match on rdfs:label:"
    )
    _add_code_block(
        d,
        """PREFIX spg: <http://spgqe.local/>
PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>
SELECT DISTINCT ?e ?label WHERE {
  ?e a spg:Entity ; rdfs:label ?label .
  FILTER(CONTAINS(LCASE(STR(?label)), "kw1") || ... || CONTAINS(..., "kwK"))
}""",
    )
    d.add_paragraph(
        "The result set B captures label-level overlap with question phrases; it complements "
        "Branch 1 when structural expansion is sparse or mis-seeded."
    )

    d.add_heading("5.4 Fusion and semantic pruning", level=2)
    d.add_paragraph(
        "Candidates are C = (A ∪ B) \\ seeds. Each c ∈ C is embedded with the same "
        "sentence-transformer used for FAISS (all-MiniLM-L6-v2). The reunion R = {question} ∪ "
        "noun-chunk probes is embedded; the score is max_{r∈R} cosine(embed(c), embed(r)). "
        "Entities with score ≥ τ are kept; seed entities are always kept. The kept set drives "
        "a natural-language graph context string appended to the dense retrieval query."
    )
    d.add_paragraph(
        "Hyperparameters n (hop depth) and τ (cosine threshold) are swept in heatmap "
        "experiments; the protocol default for primary reporting is n=2, τ=0.5. Merged-grid "
        "selection across seeds chooses a global (n*, τ*) by argmax on sample-size-weighted "
        "mean F1 (see Research Report 4)."
    )

    d.add_heading("6. Baselines and comparability", level=1)
    d.add_paragraph(
        "All pipelines share the same paragraph corpus, chunking, FAISS index (MiniLM "
        "embeddings), top-k, and Groq reader. Only retrieval conditioning differs:"
    )
    _add_table(
        d,
        ["Pipeline", "Graph use", "Description"],
        [
            ("V-RAG", "None", "FAISS on the question string only."),
            (
                "GQE-RAG (n=2)",
                "Branch 1 only",
                "n-hop neighbours appended; no τ pruning.",
            ),
            (
                "SP-GQE (n=2, τ=0.5)",
                "Branches 1+2 + pruning",
                "Primary two-branch design.",
            ),
            (
                "SP-GQE-i (n=3, τ=0.5)",
                "Iterative Branch 1",
                "Single-branch variant with per-hop pruning.",
            ),
            ("GR-RAG", "Lexical control", "Graph-adjacent entity re-ranking."),
            ("GF-RAG", "Lexical control", "Graph-adjacent entity filtering."),
        ],
    )
    d.add_paragraph(
        "No external knowledge bases are used. Comparability follows the in-repo protocol; "
        "external GraphRAG leaderboard systems are cited for positioning but not reproduced."
    )

    d.add_heading("7. Evaluation design", level=1)
    d.add_heading("7.1 Dataset and sampling", level=2)
    d.add_paragraph(
        "HotpotQA distractor development split: ten paragraphs per question (gold + "
        "distractors). Evaluation uses stratified random samples per RNG seed: approximately "
        "half bridge and half comparison questions (sample_questions). Multiple seeds are "
        "run on different days; results accumulate in results/daily_runs/ and aggregate "
        "idempotently via scripts/aggregate_daily_runs.py."
    )

    d.add_heading("7.2 Metrics", level=2)
    _add_table(
        d,
        ["Metric", "Role", "Definition"],
        [
            ("Token F1", "Primary", "Normalised token overlap vs gold short answer."),
            ("Exact match (EM)", "Secondary", "Binary normalised string equality."),
            (
                "Supporting-title recall @k",
                "Secondary",
                "Fraction of gold supporting titles in top-k chunks.",
            ),
            ("Retrieval P@k", "Diagnostic", "Fraction of top-k chunks that are supporting."),
        ],
    )
    d.add_paragraph(
        "Graph-query validity (reviewer-requested ablation): precision and recall of entity "
        "sets at four stages—Branch 1 (A), Branch 2 (B), union (A∪B), kept after τ—against "
        "spaCy-NER entities extracted from gold supporting paragraphs."
    )

    d.add_heading("7.3 Statistical analysis", level=2)
    d.add_paragraph(
        "Pipeline-level means: mean and 95% CI across seed-level means (t-interval on n seeds). "
        "Paired contrast SP-GQE − V-RAG: per (seed, qid) difference in token F1; bootstrap "
        "95% CI on the pooled mean, reported for bridge, comparison, and all subsets. "
        "Heatmap-selected (n*, τ*) reporting uses the same (n*, τ*) cell from each seed's "
        "stored grid (not per-seed argmax). Question-level paired rows at the merged cell "
        "require per-question F1 logging where not yet stored."
    )

    d.add_heading("8. Implementation", level=1)
    d.add_paragraph(
        "Repository: SP-GQE/. Core modules: src/sp_gqe/experiment/ (graph, pipelines, "
        "metrics), scripts/run_experiment.py (single-seed runs), "
        "scripts/run_heatmap_aggregate_seeds.py (multi-seed heatmap campaign), "
        "scripts/merge_heatmaps.py and scripts/aggregate_daily_runs.py (aggregation). "
        "Protocol specification: config/EXPERIMENT_PROTOCOL.md. Qualitative SPARQL traces: "
        "*__queries.md alongside daily JSON logs."
    )

    d.add_heading("References", level=1)
    for ref in [
        "Lewis, P., et al. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. NeurIPS.",
        "Yang, Z., et al. (2018). HotpotQA: A dataset for diverse, explainable multi-hop question answering. EMNLP.",
        "Izacard, G., & Grave, E. (2021). Leveraging passage retrieval with generative models for open domain QA. (Dense passage retrieval line.)",
        "Edge, D., et al. (2024). From local to global: A graph RAG approach to query-focused summarization. Microsoft technical report / blog.",
    ]:
        d.add_paragraph(ref, style="List Number")

    out = OUT_DIR / "Robert Oprescu - SP-GQE Motivation and Method.docx"
    return _save_doc(d, out)


def main() -> None:
    summary = _load_summary()
    p3 = build_report_3()
    p4 = build_report_4(summary)
    pm = build_motivation_method_docx()
    print(f"Wrote {p3}")
    print(f"Wrote {p4}")
    print(f"Wrote {pm}")


if __name__ == "__main__":
    main()
